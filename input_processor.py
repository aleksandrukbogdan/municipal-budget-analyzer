"""
Модуль обработки входных данных различных форматов.

Описание:
    Обрабатывает входные файлы любых форматов:
    - Архивы (.zip, .rar) - распаковывает
    - Excel (.xls, .xlsx) - ищет нужные листы с бюджетом
    - Word (.doc, .docx) - извлекает текст решений
    - PDF (.pdf) - извлекает текст

Основные функции:
    - process_input_folder() - обработка всей папки input/
    - find_budget_files() - поиск файлов с бюджетом расходов
    - find_decision_files() - поиск файлов решений (общие суммы)
    - extract_total_expenses() - извлечение общих расходов из решения

Автор: Автоматически сгенерировано
Версия: 1.0
"""

import os
import re
import shutil
import tempfile
import zipfile
from pathlib import Path
from dataclasses import dataclass, field
from typing import List, Dict, Optional, Tuple
import warnings
import sys

warnings.filterwarnings('ignore')

# Фикс для Windows консоли - безопасный вывод
def safe_print(msg):
    """Безопасный вывод с обработкой проблемных символов."""
    try:
        print(msg)
    except UnicodeEncodeError:
        # Заменяем проблемные символы
        print(msg.encode('cp1251', errors='replace').decode('cp1251'))

# Попытка импорта библиотек для разных форматов
try:
    import rarfile
    HAS_RARFILE = True
except ImportError:
    HAS_RARFILE = False
    print("Предупреждение: rarfile не установлен. RAR архивы не будут обрабатываться.")

try:
    from docx import Document as DocxDocument
    HAS_DOCX = True
except ImportError:
    HAS_DOCX = False
    print("Предупреждение: python-docx не установлен.")

try:
    import win32com.client
    HAS_WIN32 = True
except ImportError:
    HAS_WIN32 = False

try:
    import pdfplumber
    HAS_PDF = True
except ImportError:
    try:
        import PyPDF2
        HAS_PDF = True
    except ImportError:
        HAS_PDF = False
        print("Предупреждение: pdfplumber/PyPDF2 не установлены. PDF не будут обрабатываться.")

import pandas as pd


# --- КОНСТАНТЫ ---

INPUT_DIR = "input"
TEMP_DIR = "temp_extracted"

# Ключевые фразы для поиска нужных документов (в порядке ПРИОРИТЕТА!)
# Приоритет 1: Ведомственная структура расходов
# Приоритет 2: Распределение бюджетных ассигнований
BUDGET_EXPENSES_MARKERS_PRIORITY = [
    ("ведомственная структура расходов бюджета", 1),
    ("ведомственная структура расходов", 1),
    ("распределение бюджетных ассигнований по целевым статьям", 2),
    ("распределение бюджетных ассигнований", 2),
]

# Для обратной совместимости
BUDGET_EXPENSES_MARKERS = [m[0] for m in BUDGET_EXPENSES_MARKERS_PRIORITY]

DECISION_MARKERS = [
    "о бюджете",
    "решение думы",
    "решение совета",
    "проект решения",
    "об утверждении бюджета",
]

# Паттерны для извлечения сумм из решений
# Ищем первое вхождение суммы после ключевых слов (обычно это сумма на текущий год)
EXPENSE_PATTERNS = [
    # "Установить общий объем расходов бюджета... 1) 1 388 383 726,87 рублей на 2025 год"
    # Ивдель: "1) 1 388 383 726,87 рублей на 2025 год"
    r'объ[её]м\s+расходов[^\d]{0,150}1\)\s*([\d][\d\s]+[\d,]+)\s*руб',
    # Кировоград: "в сумме: 2 461 385 439 рублей на 2025 год"
    r'объ[её]м\s+расходов[^\d]{0,100}в\s+сумме[^\d]{0,20}([\d][\d\s]+[\d,]+)\s*руб',
    r'объ[её]м\s+расходов[^\d]{0,100}([\d][\d\s]+[\d,]+)\s*руб',
    r'расходы\s+бюджета[^\d]{0,80}([\d\s,]+[,.]?\d*)\s*(тыс|млн|руб)',
    r'расходов\s+в\s+сумме\s+([\d\s,]+[,.]?\d*)',
    r'всего\s+расходов[^\d]{0,50}([\d\s,]+[,.]?\d*)',
]

INCOME_PATTERNS = [
    # "Установить общий объем доходов бюджета... 1) 1 368 162 648,89 рублей на 2025 год"
    r'объ[её]м\s+доходов[^\d]{0,150}1\)\s*([\d][\d\s]+[\d,]+)\s*руб',
    r'объ[её]м\s+доходов[^\d]{0,100}в\s+сумме[^\d]{0,20}([\d][\d\s]+[\d,]+)\s*руб',
    r'объ[её]м\s+доходов[^\d]{0,100}([\d][\d\s]+[\d,]+)\s*руб',
    r'доходы\s+бюджета[^\d]{0,80}([\d\s,]+[,.]?\d*)\s*(тыс|млн|руб)',
    r'доходов\s+в\s+сумме\s+([\d\s,]+[,.]?\d*)',
    r'всего\s+доходов[^\d]{0,50}([\d\s,]+[,.]?\d*)',
]


@dataclass
class FileInfo:
    """Информация о найденном файле."""
    path: str
    mo_name: str  # Название муниципального образования
    file_type: str  # 'budget_expenses', 'decision', 'unknown'
    format: str  # 'excel', 'word', 'pdf'
    sheet_name: Optional[str] = None  # Для Excel - название нужного листа
    header_row: Optional[int] = None  # Для Excel - строка заголовка
    confidence: float = 0.0  # Уверенность в определении типа (0-1)
    details: Dict = field(default_factory=dict)


@dataclass 
class BudgetTotals:
    """Общие суммы бюджета из решения."""
    mo_name: str
    year: Optional[int] = None
    total_income: Optional[float] = None
    total_expenses: Optional[float] = None
    source_file: Optional[str] = None
    extraction_method: str = "not_found"


class InputProcessor:
    """
    Обработчик входных данных.
    Распаковывает архивы, находит нужные файлы, определяет их тип.
    """
    
    def __init__(self, input_dir: str = INPUT_DIR):
        self.input_dir = input_dir
        self.temp_dir = TEMP_DIR
        self.found_files: List[FileInfo] = []
        self.budget_totals: Dict[str, BudgetTotals] = {}  # МО -> общие суммы
        
    def process_all(self) -> Tuple[List[FileInfo], Dict[str, BudgetTotals]]:
        """
        Главный метод: обрабатывает всю папку input/.
        
        Returns:
            Tuple (список найденных файлов с бюджетом, словарь общих сумм по МО)
        """
        safe_print(f"\n{'='*60}")
        safe_print("ОБРАБОТКА ВХОДНЫХ ДАННЫХ")
        safe_print(f"{'='*60}")
        
        if not os.path.exists(self.input_dir):
            safe_print(f"Папка {self.input_dir} не найдена!")
            return [], {}
        
        # Создаём временную папку для распакованных файлов
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)
        os.makedirs(self.temp_dir, exist_ok=True)
        
        try:
            # 1. Сканируем и распаковываем архивы
            self._scan_and_extract(self.input_dir)
            
            # 2. Классифицируем все найденные файлы
            self._classify_files()
            
            # 3. Извлекаем общие суммы из решений
            self._extract_totals_from_decisions()
            
            # Выводим результаты
            self._print_summary()
            
        finally:
            # Очищаем временную папку (опционально)
            # shutil.rmtree(self.temp_dir, ignore_errors=True)
            pass
        
        return self.found_files, self.budget_totals
    
    def _scan_and_extract(self, directory: str):
        """Рекурсивно сканирует папку, распаковывает архивы."""
        safe_print(f"\nСканирование: {directory}")
        
        for item in os.listdir(directory):
            item_path = os.path.join(directory, item)
            
            if os.path.isdir(item_path):
                # Рекурсивно обрабатываем подпапки
                self._scan_and_extract(item_path)
                
            elif os.path.isfile(item_path):
                ext = os.path.splitext(item)[1].lower()
                
                # Определяем МО по имени папки или файла
                mo_name = self._extract_mo_name(item_path)
                
                if ext == '.zip':
                    self._extract_zip(item_path, mo_name)
                elif ext == '.rar' and HAS_RARFILE:
                    self._extract_rar(item_path, mo_name)
                elif ext in ['.xls', '.xlsx', '.doc', '.docx', '.pdf']:
                    # Добавляем файл для дальнейшей классификации
                    if not item.startswith('~$'):  # Пропускаем временные файлы
                        self.found_files.append(FileInfo(
                            path=item_path,
                            mo_name=mo_name,
                            file_type='unknown',
                            format=self._get_format(ext)
                        ))
    
    def _extract_zip(self, zip_path: str, mo_name: str):
        """Распаковывает ZIP архив."""
        safe_print(f"  Распаковка ZIP: {os.path.basename(zip_path)}")
        
        extract_to = os.path.join(self.temp_dir, mo_name, os.path.splitext(os.path.basename(zip_path))[0])
        os.makedirs(extract_to, exist_ok=True)
        
        try:
            with zipfile.ZipFile(zip_path, 'r') as zf:
                zf.extractall(extract_to)
            # Рекурсивно обрабатываем распакованные файлы
            self._scan_and_extract(extract_to)
        except Exception as e:
            safe_print(f"    Ошибка распаковки: {e}")
    
    def _extract_rar(self, rar_path: str, mo_name: str):
        """Распаковывает RAR архив."""
        safe_print(f"  Распаковка RAR: {os.path.basename(rar_path)}")
        
        extract_to = os.path.join(self.temp_dir, mo_name, os.path.splitext(os.path.basename(rar_path))[0])
        os.makedirs(extract_to, exist_ok=True)
        
        try:
            with rarfile.RarFile(rar_path, 'r') as rf:
                rf.extractall(extract_to)
            # Рекурсивно обрабатываем распакованные файлы
            self._scan_and_extract(extract_to)
        except Exception as e:
            safe_print(f"    Ошибка распаковки: {e}")
    
    def _extract_mo_name(self, filepath: str) -> str:
        """Извлекает название МО из пути к файлу."""
        parts = Path(filepath).parts
        
        # Ищем папку с названием МО (обычно содержит "_Бюджет")
        for part in parts:
            if "_бюджет" in part.lower() or "_budget" in part.lower():
                return part.split('_')[0]
        
        # Или берём из имени файла
        filename = os.path.basename(filepath)
        if '_' in filename:
            return filename.split('_')[0]
        
        # Или берём имя папки
        parent = os.path.basename(os.path.dirname(filepath))
        if parent != 'input' and parent != self.temp_dir:
            return parent.split('_')[0]
        
        return "Unknown"
    
    def _get_format(self, ext: str) -> str:
        """Определяет формат файла по расширению."""
        if ext in ['.xls', '.xlsx']:
            return 'excel'
        elif ext in ['.doc', '.docx']:
            return 'word'
        elif ext == '.pdf':
            return 'pdf'
        return 'unknown'
    
    def _classify_files(self):
        """Классифицирует все найденные файлы."""
        safe_print(f"\nКлассификация файлов ({len(self.found_files)} шт.)...")
        
        for file_info in self.found_files:
            if file_info.format == 'excel':
                self._classify_excel(file_info)
            elif file_info.format == 'word':
                self._classify_word(file_info)
            elif file_info.format == 'pdf':
                self._classify_pdf(file_info)
    
    def _classify_excel(self, file_info: FileInfo):
        """Классифицирует Excel файл, проверяя содержимое листов."""
        try:
            xls = pd.ExcelFile(file_info.path)
            
            for sheet_name in xls.sheet_names:
                # Читаем первые 30 строк листа
                try:
                    df = pd.read_excel(file_info.path, sheet_name=sheet_name, 
                                       header=None, nrows=30)
                except:
                    continue
                
                # Собираем весь текст из первых строк
                text_content = ""
                for idx, row in df.iterrows():
                    for val in row:
                        if pd.notna(val):
                            text_content += " " + str(val).lower()
                
                # Проверяем на маркеры бюджета расходов (с приоритетом!)
                for marker, priority in BUDGET_EXPENSES_MARKERS_PRIORITY:
                    if marker in text_content:
                        file_info.file_type = 'budget_expenses'
                        file_info.sheet_name = sheet_name
                        file_info.confidence = 0.9
                        file_info.details['marker_priority'] = priority
                        file_info.details['marker_found'] = marker
                        
                        # Ищем строку заголовка
                        for idx, row in df.iterrows():
                            row_text = " ".join([str(x).lower() for x in row if pd.notna(x)])
                            if "наименование" in row_text and ("раздел" in row_text or "код" in row_text):
                                file_info.header_row = idx
                                break
                        
                        safe_print(f"  [OK] BUDGET_EXPENSES (P{priority}): {os.path.basename(file_info.path)} (лист: {sheet_name})")
                        return
                
                # Проверяем на маркеры решения (редко в Excel, но возможно)
                for marker in DECISION_MARKERS:
                    if marker in text_content:
                        file_info.file_type = 'decision'
                        file_info.sheet_name = sheet_name
                        file_info.confidence = 0.7
                        safe_print(f"  [?] DECISION (Excel): {os.path.basename(file_info.path)}")
                        return
            
            # Если не нашли маркеры, проверяем имя файла
            filename_lower = os.path.basename(file_info.path).lower()
            if "прил" in filename_lower and any(n in filename_lower for n in ['4', '5', '9', '10']):
                file_info.file_type = 'budget_expenses'
                file_info.sheet_name = xls.sheet_names[0]
                file_info.confidence = 0.6
                safe_print(f"  [?] BUDGET_EXPENSES (по имени): {os.path.basename(file_info.path)}")
            
        except Exception as e:
            safe_print(f"  [ERR] Ошибка чтения {os.path.basename(file_info.path)}: {e}")
    
    def _classify_word(self, file_info: FileInfo):
        """Классифицирует Word файл."""
        # Извлекаем текст и таблицы
        result = self._extract_text_from_word(file_info.path, with_tables=True)
        if isinstance(result, tuple):
            text, tables = result
        else:
            text, tables = result, []
        
        if not text:
            return
        
        text_lower = text.lower()
        filename_lower = os.path.basename(file_info.path).lower()
        
        # Определяем тип файла по имени
        is_appendix = "прил" in filename_lower
        is_decision_by_name = ("решени" in filename_lower or "проект" in filename_lower) and "прил" not in filename_lower
        
        # ВАЖНО: Проверяем на типичные маркеры решения в НАЧАЛЕ документа
        # Если документ начинается с "РЕШЕНИЕ" или содержит "ДУМА...РЕШИЛА" - это DECISION
        text_start = text[:2000].lower()
        is_decision_by_content = (
            "дума" in text_start and "решила" in text_start or
            "решение №" in text_start or
            "решение  №" in text_start or  # С двойным пробелом
            text_start.strip().startswith("решение") or
            "об утверждении бюджета" in text_start or
            "о бюджете" in text_start and "на 20" in text_start and "год" in text_start
        )
        
        # Если файл - решение по содержимому начала, это DECISION
        if is_decision_by_content and not is_appendix:
            file_info.file_type = 'decision'
            file_info.confidence = 0.95
            file_info.details['text_preview'] = text[:500]
            safe_print(f"  [OK] DECISION (по содержимому): {os.path.basename(file_info.path)}")
            return
        
        # Если файл - решение по имени, сначала проверяем на DECISION
        if is_decision_by_name:
            for marker in DECISION_MARKERS:
                if marker in text_lower:
                    file_info.file_type = 'decision'
                    file_info.confidence = 0.95
                    file_info.details['text_preview'] = text[:500]
                    safe_print(f"  [OK] DECISION: {os.path.basename(file_info.path)}")
                    return
        
        # Проверяем маркеры бюджета расходов (с приоритетом!)
        for marker, priority in BUDGET_EXPENSES_MARKERS_PRIORITY:
            if marker in text_lower:
                file_info.file_type = 'budget_expenses'
                file_info.confidence = 0.9
                file_info.details['has_tables'] = len(tables) > 0
                file_info.details['tables_count'] = len(tables)
                file_info.details['marker_priority'] = priority
                file_info.details['marker_found'] = marker
                # Сохраняем таблицы для последующей обработки в extract_articles
                if tables:
                    file_info.details['tables_data'] = self._convert_tables_to_dataframes(tables)
                safe_print(f"  [OK] BUDGET_EXPENSES (Word, P{priority}): {os.path.basename(file_info.path)}")
                return
        
        # Проверяем наличие таблицы с бюджетными данными (для приложений)
        if is_appendix and tables:
            # Проверяем первую таблицу на наличие бюджетных колонок
            for table in tables[:3]:  # Проверяем первые 3 таблицы
                if len(table) > 5:  # Достаточно строк
                    header_text = " ".join([str(cell).lower() for row in table[:3] for cell in row])
                    if ("код" in header_text and ("раздел" in header_text or "целев" in header_text)) or \
                       ("наименование" in header_text and "сумм" in header_text):
                        file_info.file_type = 'budget_expenses'
                        file_info.confidence = 0.85
                        file_info.details['has_tables'] = True
                        file_info.details['tables_count'] = len(tables)
                        file_info.details['source'] = 'word_table'
                        file_info.details['marker_priority'] = 2  # Приоритет 2 для Word таблиц без явного маркера
                        # Сохраняем таблицы для последующей обработки
                        file_info.details['tables_data'] = self._convert_tables_to_dataframes(tables)
                        safe_print(f"  [OK] BUDGET_EXPENSES (Word таблица): {os.path.basename(file_info.path)}")
                        return
        
        # Проверяем маркеры решения (если ещё не определили как DECISION выше)
        if not is_decision_by_name:
            for marker in DECISION_MARKERS:
                if marker in text_lower:
                    # Дополнительная проверка - не приложение ли это
                    if not is_appendix:
                        file_info.file_type = 'decision'
                        file_info.confidence = 0.9
                        file_info.details['text_preview'] = text[:500]
                        safe_print(f"  [OK] DECISION: {os.path.basename(file_info.path)}")
                        return
        
        # Проверяем имя файла на решение (если не приложение)
        if not is_appendix and ("бюджет" in filename_lower):
            file_info.file_type = 'decision'
            file_info.confidence = 0.6
            file_info.details['text_preview'] = text[:500]
            safe_print(f"  [?] DECISION (по имени): {os.path.basename(file_info.path)}")
            return
        
        # Проверяем имя файла на приложение с расходами
        if is_appendix and any(n in filename_lower for n in ['4', '5', '3']):
            file_info.file_type = 'budget_expenses'
            file_info.confidence = 0.5
            file_info.details['has_tables'] = len(tables) > 0
            safe_print(f"  [?] BUDGET_EXPENSES (по имени прил): {os.path.basename(file_info.path)}")
    
    def _classify_pdf(self, file_info: FileInfo):
        """Классифицирует PDF файл."""
        text = self._extract_text_from_pdf(file_info.path)
        if not text:
            return
        
        text_lower = text.lower()
        
        # Аналогично Word
        for marker in DECISION_MARKERS:
            if marker in text_lower:
                file_info.file_type = 'decision'
                file_info.confidence = 0.9
                file_info.details['text_preview'] = text[:500]
                safe_print(f"  [OK] DECISION (PDF): {os.path.basename(file_info.path)}")
                return
        
        for marker in BUDGET_EXPENSES_MARKERS:
            if marker in text_lower:
                file_info.file_type = 'budget_expenses'
                file_info.confidence = 0.8
                safe_print(f"  [OK] BUDGET_EXPENSES (PDF): {os.path.basename(file_info.path)}")
                return
    
    def _convert_tables_to_dataframes(self, tables: List[List[List[str]]]) -> List:
        """
        Конвертирует таблицы из списков в DataFrame.
        
        Args:
            tables: Список таблиц в формате [[[cell, cell], [cell, cell]], ...]
        
        Returns:
            Список DataFrame
        """
        import pandas as pd
        result = []
        for table in tables:
            if table and len(table) > 0:
                try:
                    df = pd.DataFrame(table)
                    result.append(df)
                except Exception:
                    pass
        return result
    
    def _extract_text_from_word(self, filepath: str, with_tables: bool = False) -> tuple:
        """
        Извлекает текст из Word документа.
        
        Args:
            filepath: Путь к файлу
            with_tables: Также извлекать таблицы
        
        Returns:
            Если with_tables=False: строка с текстом
            Если with_tables=True: (текст, список таблиц)
        """
        ext = os.path.splitext(filepath)[1].lower()
        text = ""
        tables = []
        
        # Для .docx используем python-docx
        if ext == '.docx' and HAS_DOCX:
            try:
                doc = DocxDocument(filepath)
                text = "\n".join([para.text for para in doc.paragraphs])
                
                if with_tables:
                    for table in doc.tables:
                        table_data = []
                        for row in table.rows:
                            row_data = [cell.text.strip() for cell in row.cells]
                            table_data.append(row_data)
                        tables.append(table_data)
            except:
                pass
        
        # Для .doc используем win32com (только Windows)
        if not text and HAS_WIN32:
            word = None
            doc = None
            try:
                word = win32com.client.Dispatch("Word.Application")
                word.Visible = False
                word.DisplayAlerts = 0  # Отключаем диалоги
                
                doc = word.Documents.Open(
                    os.path.abspath(filepath),
                    ReadOnly=True,
                    AddToRecentFiles=False
                )
                text = doc.Content.Text
                
                if with_tables and doc.Tables.Count > 0:
                    for i in range(1, min(doc.Tables.Count + 1, 10)):  # Макс 10 таблиц
                        try:
                            table = doc.Tables(i)
                            table_data = []
                            for row_idx in range(1, min(table.Rows.Count + 1, 500)):  # Макс 500 строк
                                row_data = []
                                for col_idx in range(1, table.Columns.Count + 1):
                                    try:
                                        cell_text = table.Cell(row_idx, col_idx).Range.Text
                                        cell_text = cell_text.strip().replace('\r\x07', '').replace('\r', '')
                                        row_data.append(cell_text)
                                    except:
                                        row_data.append('')
                                table_data.append(row_data)
                            tables.append(table_data)
                        except:
                            continue
                
            except Exception as e:
                pass
            finally:
                try:
                    if doc:
                        doc.Close(False)
                except:
                    pass
                try:
                    if word:
                        word.Quit()
                except:
                    pass
        
        # Fallback: попробовать прочитать как текст (для простых случаев)
        if not text:
            try:
                with open(filepath, 'rb') as f:
                    content = f.read()
                    text = content.decode('utf-8', errors='ignore')
                    text = re.sub(r'[^\x20-\x7E\u0400-\u04FF\s]', ' ', text)
            except:
                pass
        
        if with_tables:
            return text, tables
        return text
    
    def _extract_text_from_pdf(self, filepath: str) -> str:
        """Извлекает текст из PDF."""
        try:
            import pdfplumber
            with pdfplumber.open(filepath) as pdf:
                text = ""
                for page in pdf.pages[:10]:  # Первые 10 страниц
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
                return text
        except:
            pass
        
        try:
            import PyPDF2
            with open(filepath, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                text = ""
                for page in reader.pages[:10]:
                    text += page.extract_text() + "\n"
                return text
        except:
            pass
        
        return ""
    
    def _extract_totals_from_decisions(self):
        """Извлекает общие суммы доходов/расходов из файлов решений."""
        safe_print(f"\nИзвлечение общих сумм из решений...")
        
        decision_files = [f for f in self.found_files if f.file_type == 'decision']
        
        for file_info in decision_files:
            # Извлекаем текст
            if file_info.format == 'word':
                result = self._extract_text_from_word(file_info.path, with_tables=False)
                text = result if isinstance(result, str) else result[0] if result else ""
            elif file_info.format == 'pdf':
                text = self._extract_text_from_pdf(file_info.path)
            else:
                continue
            
            if not text:
                continue
            
            # Нормализуем текст
            text = text.replace('\n', ' ').replace('\r', ' ')
            text = re.sub(r'\s+', ' ', text)
            
            # Извлекаем год
            year_match = re.search(r'на\s+(20[2-3]\d)\s*год', text, re.IGNORECASE)
            year = int(year_match.group(1)) if year_match else None
            
            # Извлекаем расходы
            expenses = self._extract_amount(text, EXPENSE_PATTERNS)
            income = self._extract_amount(text, INCOME_PATTERNS)
            
            mo_name = file_info.mo_name
            
            if expenses or income:
                self.budget_totals[mo_name] = BudgetTotals(
                    mo_name=mo_name,
                    year=year,
                    total_income=income,
                    total_expenses=expenses,
                    source_file=os.path.basename(file_info.path),
                    extraction_method='regex'
                )
                safe_print(f"  [OK] {mo_name}: доходы={income}, расходы={expenses} (год: {year})")
            else:
                safe_print(f"  [--] {mo_name}: суммы не найдены в {os.path.basename(file_info.path)}")
    
    def _extract_amount(self, text: str, patterns: List[str]) -> Optional[float]:
        """Извлекает сумму из текста по паттернам."""
        for pattern in patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                num_str = match.group(1)
                # Очищаем число: убираем пробелы, заменяем запятую на точку
                num_str = num_str.replace(' ', '')
                # Если есть запятая - это десятичный разделитель (российский формат)
                # "1368162648,89" -> "1368162648.89"
                if ',' in num_str:
                    # Заменяем последнюю запятую на точку
                    parts = num_str.rsplit(',', 1)
                    if len(parts) == 2 and len(parts[1]) <= 2:  # Десятичная часть
                        num_str = parts[0] + '.' + parts[1]
                    else:
                        num_str = num_str.replace(',', '')  # Это разделитель тысяч
                
                try:
                    value = float(num_str)
                    # Проверяем единицу измерения
                    if len(match.groups()) > 1:
                        unit = match.group(2).lower() if match.group(2) else ""
                        if 'млн' in unit:
                            value *= 1000  # Переводим в тыс.
                    return value
                except ValueError:
                    continue
        return None
    
    def _print_summary(self):
        """Выводит сводку найденных файлов."""
        safe_print(f"\n{'='*60}")
        safe_print("ИТОГИ СКАНИРОВАНИЯ")
        safe_print(f"{'='*60}")
        
        budget_files = [f for f in self.found_files if f.file_type == 'budget_expenses']
        decision_files = [f for f in self.found_files if f.file_type == 'decision']
        unknown_files = [f for f in self.found_files if f.file_type == 'unknown']
        
        safe_print(f"\nФайлы с расходами бюджета: {len(budget_files)}")
        for f in budget_files:
            sheet_info = f" (лист: {f.sheet_name})" if f.sheet_name else ""
            safe_print(f"  - {f.mo_name}: {os.path.basename(f.path)}{sheet_info}")
        
        safe_print(f"\nФайлы решений: {len(decision_files)}")
        for f in decision_files:
            safe_print(f"  - {f.mo_name}: {os.path.basename(f.path)}")
        
        if unknown_files:
            safe_print(f"\nНеклассифицированные файлы: {len(unknown_files)}")
            for f in unknown_files:
                safe_print(f"  - {os.path.basename(f.path)}")
        
        safe_print(f"\nОбщие суммы бюджетов:")
        for mo, totals in self.budget_totals.items():
            safe_print(f"  - {mo}: доходы={totals.total_income}, расходы={totals.total_expenses}")
    
    def get_budget_expense_files(self, one_per_mo: bool = True) -> List[FileInfo]:
        """
        Возвращает файлы с расходами бюджета.
        
        Args:
            one_per_mo: Если True - возвращает только ОДИН лучший файл для каждого МО
                        по приоритету маркеров (Ведомственная структура > Распределение ассигнований)
        """
        all_budget_files = [f for f in self.found_files if f.file_type == 'budget_expenses']
        
        if not one_per_mo:
            return all_budget_files
        
        # Выбираем лучший файл для каждого МО
        best_by_mo = {}
        
        for f in all_budget_files:
            mo = f.mo_name
            priority = f.details.get('marker_priority', 99)  # 99 = низкий приоритет
            
            if mo not in best_by_mo:
                best_by_mo[mo] = f
            else:
                current_best = best_by_mo[mo]
                current_priority = current_best.details.get('marker_priority', 99)
                
                # Меньший номер приоритета = лучше (1 лучше чем 2)
                if priority < current_priority:
                    safe_print(f"  [INFO] {mo}: заменяем {os.path.basename(current_best.path)} на {os.path.basename(f.path)} (приоритет {priority} < {current_priority})")
                    best_by_mo[mo] = f
                elif priority == current_priority:
                    # При равном приоритете предпочитаем Excel над Word
                    if f.format == 'excel' and current_best.format == 'word':
                        safe_print(f"  [INFO] {mo}: предпочитаем Excel {os.path.basename(f.path)}")
                        best_by_mo[mo] = f
        
        result = list(best_by_mo.values())
        
        if len(result) < len(all_budget_files):
            safe_print(f"\n  Выбрано {len(result)} файлов бюджета (по одному на МО) из {len(all_budget_files)} найденных")
        
        return result
    
    def get_decision_files(self) -> List[FileInfo]:
        """Возвращает только файлы решений."""
        return [f for f in self.found_files if f.file_type == 'decision']


def process_input_folder(input_dir: str = INPUT_DIR):
    """
    Удобная функция для запуска обработки.
    
    Returns:
        Tuple (список файлов бюджета, словарь общих сумм по МО)
    """
    processor = InputProcessor(input_dir)
    return processor.process_all()


if __name__ == "__main__":
    # Тестовый запуск
    files, totals = process_input_folder()
    
    print("\n" + "="*60)
    print("ГОТОВО К ПЕРЕДАЧЕ В EXTRACT_ARTICLES.PY")
    print("="*60)
    
    for f in files:
        if f.file_type == 'budget_expenses':
            safe_print(f"\n{f.mo_name}:")
            safe_print(f"  Файл: {f.path}")
            safe_print(f"  Лист: {f.sheet_name}")
            safe_print(f"  Строка заголовка: {f.header_row}")

